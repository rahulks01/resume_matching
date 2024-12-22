import streamlit as st
import pandas as pd
import requests
import io  # Importing the io module for StringIO
import os
import fitz  # PyMuPDF for PDF extraction
import docx  # python-docx for DOCX extraction
import re
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import gdown
# Load the model once and cache it
@st.cache_resource
def load_model():
    return SentenceTransformer('all-MiniLM-L6-v2')

model = load_model()

@st.cache_data
def load_and_cache_job_data():
    url = 'https://drive.google.com/uc?id=1w6BDl80dBlRt9LuycrucswrBfD7xfM_R'
    
    try:
        output = 'dataset.csv'
        gdown.download(url, output, quiet=False)
        df_jobs = pd.read_csv(output)
        
        job_descriptions = df_jobs['Description'].fillna('').tolist()
        job_titles = df_jobs['Title'].fillna('Unknown').tolist()
        job_vectors = model.encode(job_descriptions, batch_size=32, show_progress_bar=True)
        
        return job_descriptions, job_titles, job_vectors
    except Exception as e:
        st.error(f"Error loading dataset: {e}")
        return [], [], []

# Load data
job_descriptions, job_titles, job_vectors = load_and_cache_job_data()


job_descriptions, job_titles, job_vectors = load_and_cache_job_data()

def extract_text_from_pdf(pdf_path):
    try:
        doc = fitz.open(pdf_path)
        text = ""
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            text += page.get_text("text")
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return None

def extract_text_from_docx(docx_path):
    try:
        doc = docx.Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        st.error(f"Error reading DOCX: {e}")
        return None

def extract_text_from_resume(uploaded_file):
    ext = os.path.splitext(uploaded_file.name)[1].lower() 
    
    if ext == '.pdf':
        try:
            doc = fitz.open(stream=uploaded_file.read(), filetype='pdf')
            text = ""
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text += page.get_text("text")
            return text
        except Exception as e:
            st.error(f"Error reading PDF: {e}")
            return None
    
    elif ext == '.docx':
        try:
            doc = docx.Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            st.error(f"Error reading DOCX: {e}")
            return None
    
    else:
        st.error(f"Unsupported file format: {ext}")
        return None

def preprocess_text(text):
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'[^\w\s]', '', text)
    return text.strip()

def predict_jobs_for_resume(resume_text):
    cleaned_resume = preprocess_text(resume_text)
    resume_vector = model.encode([cleaned_resume])
    similarities = cosine_similarity(resume_vector, job_vectors)[0]
    top_indices = similarities.argsort()[-5:][::-1]
    top_jobs = [(job_titles[i], job_descriptions[i], similarities[i]) for i in top_indices]
    return top_jobs

# Streamlit UI
st.title("Resume-Based Job Matching")

uploaded_file = st.file_uploader("Upload your resume (PDF or DOCX)", type=['pdf', 'docx'])

if uploaded_file:
    with st.spinner("Extracting and processing the resume..."):
        resume_text = extract_text_from_resume(uploaded_file)
        if resume_text:
            st.write("Resume text extracted successfully.")
            with st.spinner("Matching jobs..."):
                top_jobs = predict_jobs_for_resume(resume_text)
                st.success("Job matching completed!")
                
                # Display results
                st.subheader("Top 5 Job Matches")
                for i, (title, description, score) in enumerate(top_jobs, start=1):
                    st.markdown(f"### {i}. {title} (Score: {score:.2f})")
                    st.write(description[:200] + "...")  # Show first 200 chars of description


